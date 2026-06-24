import sys
import os
import json
import base64
import shutil
import sqlite3
import logging
import hashlib
import datetime
import uuid
import webview

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Definición de rutas base
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def get_writable_base():
    """Retorna una ruta segura para escribir datos (evita bloqueos en Program Files)."""
    # Si estamos en modo desarrollo o portable (carpeta escribible), usar BASE_DIR
    if os.access(BASE_DIR, os.W_OK):
        return BASE_DIR
    
    # Si es una instalación en Program Files, usar LocalAppData del usuario
    app_data = os.getenv('LOCALAPPDATA')
    if app_data:
        path = os.path.join(app_data, 'JURIS-GESTIÓN-PRO')
        os.makedirs(path, exist_ok=True)
        return path
    return BASE_DIR

WRITABLE_DIR = get_writable_base()

# Rutas de Lectura (Pueden estar en Program Files)
UI_PATH = os.path.join(BASE_DIR, 'ui', 'juris_gestion_pro_app.html')

# Rutas de Escritura (Deben estar en WRITABLE_DIR)
CONFIG_PATH = os.path.join(WRITABLE_DIR, 'config', 'incidentes_procesales.json')
DB_PATH = os.path.join(WRITABLE_DIR, 'data', 'juris_gestion_pro.db')

# Asegurar que existan las carpetas de escritura
os.makedirs(os.path.join(WRITABLE_DIR, 'config'), exist_ok=True)
os.makedirs(os.path.join(WRITABLE_DIR, 'data'), exist_ok=True)
os.makedirs(os.path.join(WRITABLE_DIR, 'assets'), exist_ok=True)

# Si el config no existe en la carpeta escribible, copiar el original desde BASE_DIR
ORIGINAL_CONFIG = os.path.join(BASE_DIR, 'config', 'incidentes_procesales.json')
if not os.path.exists(CONFIG_PATH) and os.path.exists(ORIGINAL_CONFIG):
    try:
        shutil.copy2(ORIGINAL_CONFIG, CONFIG_PATH)
    except Exception: pass

# Garantizar que los módulos src/ sean importables desde cualquier contexto
if BASE_DIR not in sys.path:
    sys.path.append(BASE_DIR)


def _get_db_connection():
    """Abre una conexión SQLite con row_factory configurado."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


import time
import functools

def log_performance(func):
    """Decorador para medir y registrar el tiempo de ejecución de las llamadas API."""
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.perf_counter()
        result = func(*args, **kwargs)
        end_time = time.perf_counter()
        duration = (end_time - start_time) * 1000
        logging.info(f"⚡ [API-BACKEND] {func.__name__} ejecutada en {duration:.2f}ms")
        return result
    return wrapper


# ══════════════════════════════════════════════════════════════
# HELPERS DE LICENCIAS (nivel módulo)
# ══════════════════════════════════════════════════════════════
_LICENSE_SECRET = "JURIS_PRO_FER_ARDON_2026"
_BASE_DATE = datetime.date(2024, 1, 1)

def _get_hardware_id():
    """
    Genera un identificador único de hardware de 10 caracteres
    basado en la dirección MAC de la interfaz de red principal.
    """
    try:
        mac_int = uuid.getnode()
        # Convertir el entero MAC a hex de 12 dígitos, luego tomar 10
        mac_hex = f"{mac_int:012X}"
        return mac_hex[:10]
    except Exception:
        return "UNIVERSAL"


def _validar_key_interna(key, hw_id):
    """
    Valida una clave de licencia contra el Hardware ID del equipo.
    La clave puede venir con o sin guiones (formato XXXXX-XXXXX-XXXXX-XXXXX).
    Retorna dict con: valida (bool), mensaje (str).
    """
    try:
        # Limpiar y normalizar
        key_limpia = key.strip().upper().replace('-', '')
        if len(key_limpia) != 20:
            return {'valida': False, 'mensaje': 'La clave debe tener 20 caracteres.'}

        # Extraer componentes
        exp_hex   = key_limpia[:4]
        signature = key_limpia[4:20]

        # Re-calcular la firma esperada con el HW ID local
        hw_id_norm = hw_id.strip().upper()
        firma_esperada = hashlib.sha256(
            f"{exp_hex}{hw_id_norm}{_LICENSE_SECRET}".encode()
        ).hexdigest()[:16].upper()

        # También aceptar clave UNIVERSAL (hw_id = "UNIVERSAL")
        firma_universal = hashlib.sha256(
            f"{exp_hex}UNIVERSAL{_LICENSE_SECRET}".encode()
        ).hexdigest()[:16].upper()

        if signature != firma_esperada and signature != firma_universal:
            return {'valida': False, 'mensaje': 'Clave inválida. Verifique la clave o el Hardware ID.'}

        # Verificar expiración
        if exp_hex == 'FFFF':
            return {'valida': True, 'mensaje': 'Licencia PERMANENTE activa. ✅'}

        dias_totales = int(exp_hex, 16)
        fecha_exp = _BASE_DATE + datetime.timedelta(days=dias_totales)
        hoy = datetime.date.today()

        if hoy > fecha_exp:
            return {'valida': False, 'mensaje': f'Licencia vencida el {fecha_exp.strftime("%d/%m/%Y")}. Contacte al proveedor para renovar.'}

        dias_restantes = (fecha_exp - hoy).days
        return {'valida': True, 'mensaje': f'Licencia activa. Vence el {fecha_exp.strftime("%d/%m/%Y")} ({dias_restantes} días restantes). ✅'}

    except Exception as e:
        logging.error(f"Error al validar licencia: {e}")
        return {'valida': False, 'mensaje': f'Error interno al validar la clave: {e}'}


class JurisAPI:
    """Clase puente entre Python (Backend) y HTML/JS (Frontend)"""
    def __init__(self):
        from src.database import inicializar_db
        inicializar_db()
        self._migrar_columnas()

    def _migrar_columnas(self):
        """Aplica migraciones de columnas adicionales de forma segura al inicio."""
        migraciones = [
            "ALTER TABLE configuracion ADD COLUMN logo_path TEXT DEFAULT '../assets/juris_gestion_pro_logo.png'",
            "ALTER TABLE configuracion ADD COLUMN pin TEXT DEFAULT ''",
            "ALTER TABLE configuracion ADD COLUMN incidentes_path TEXT DEFAULT ''",
            "ALTER TABLE actuaciones ADD COLUMN monto_involucrado REAL DEFAULT 0",
            "ALTER TABLE clientes ADD COLUMN correo TEXT",
            "ALTER TABLE configuracion ADD COLUMN licencia_key TEXT DEFAULT ''",
        ]
        try:
            conn = _get_db_connection()
            for sql in migraciones:
                try:
                    conn.execute(sql)
                except sqlite3.OperationalError:
                    pass  # Columna ya existe
            conn.commit()
            conn.close()
        except Exception as e:
            logging.warning(f"Migracion de columnas: {e}")

    def _get_ruta_incidentes(self, cursor):
        """Resuelve la ruta del archivo de incidentes desde configuración."""
        try:
            cursor.execute('SELECT incidentes_path FROM configuracion WHERE id = 1')
            row = cursor.fetchone()
            ruta = row['incidentes_path'] if row and row['incidentes_path'] else ''
            return ruta if ruta else CONFIG_PATH
        except Exception:
            return CONFIG_PATH

    # ══════════════════════════════════════════════════════════════
    # SISTEMA DE LICENCIAS
    # ══════════════════════════════════════════════════════════════

    def get_estado_licencia(self):
        """
        Verifica si la aplicación tiene una licencia válida.
        Retorna un dict con: activado (bool), hw_id (str), info (str).
        """
        hw_id = _get_hardware_id()
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('SELECT licencia_key FROM configuracion WHERE id = 1')
            row = cursor.fetchone()
            conn.close()
            key_guardada = row['licencia_key'] if row and row['licencia_key'] else ''
        except Exception:
            key_guardada = ''

        if not key_guardada:
            return {'activado': False, 'hw_id': hw_id, 'info': 'Sin licencia registrada.'}

        resultado = _validar_key_interna(key_guardada, hw_id)
        return {'activado': resultado['valida'], 'hw_id': hw_id, 'info': resultado['mensaje']}

    def validar_licencia_key(self, key):
        """
        Recibe una clave de 20 caracteres (con o sin guiones),
        la valida contra el Hardware ID del equipo y la guarda si es correcta.
        Retorna dict con: valida (bool), mensaje (str).
        """
        hw_id = _get_hardware_id()
        resultado = _validar_key_interna(key, hw_id)
        if resultado['valida']:
            try:
                conn = _get_db_connection()
                key_limpia = key.strip().upper().replace('-', '')
                key_fmt = f"{key_limpia[:5]}-{key_limpia[5:10]}-{key_limpia[10:15]}-{key_limpia[15:20]}"
                conn.execute("UPDATE configuracion SET licencia_key = ? WHERE id = 1", (key_fmt,))
                conn.commit()
                conn.close()
                logging.info("Licencia válida registrada exitosamente.")
            except Exception as e:
                logging.error(f"Error al guardar licencia: {e}")
                return {'valida': False, 'mensaje': f'Error al guardar la licencia: {e}'}
        return resultado

    @log_performance
    def get_incidentes(self):
        """Lee y retorna el catálogo de incidentes desde el archivo JSON (interno o externo)."""
        try:
            conn = _get_db_connection()
            ruta_incidentes = self._get_ruta_incidentes(conn.cursor())
            conn.close()
            with open(ruta_incidentes, 'r', encoding='utf-8') as f:
                incidentes = json.load(f)
            logging.info(f"Se cargaron {len(incidentes)} incidentes desde: {ruta_incidentes}")
            return incidentes
        except Exception as e:
            logging.error(f"Error al leer incidentes: {e}")
            return [{"error": str(e)}]
            
    @log_performance
    def get_resumen_casos(self):
        """Consulta los casos por estado desde SQLite para los contadores."""
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('SELECT estado, COUNT(*) FROM expedientes GROUP BY estado')
            resultados = dict(cursor.fetchall())
            conn.close()
            return {
                "Activo": resultados.get("Activo", 0),
                "En Proceso": resultados.get("En Proceso", 0),
                "Cerrado": resultados.get("Cerrado", 0),
            }
        except Exception as e:
            logging.error(f"Error al contar casos: {e}")
            return {"Activo": 0, "En Proceso": 0, "Cerrado": 0}

    @log_performance
    def get_casos(self, estado="Activo"):
        """Consulta los casos por estado desde SQLite para la tabla."""
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                SELECT e.id_expediente, c.nombre_completo, e.estado, e.fecha_creacion, e.juzgado
                FROM expedientes e
                JOIN clientes c ON e.id_cliente = c.id_cliente
                WHERE e.estado = ?
            ''', (estado,))
            casos = [dict(row) for row in cursor.fetchall()]
            conn.close()
            return casos
        except Exception as e:
            logging.error(f"Error al leer casos: {e}")
            return []

    def get_protocolos(self):
        """Consulta el libro de protocolo notarial desde SQLite."""
        return self.get_protocolo()

    @log_performance
    def get_clientes(self):
        """Consulta el directorio de clientes (CRM)."""
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM clientes ORDER BY id_cliente DESC')
            clientes = [dict(row) for row in cursor.fetchall()]
            conn.close()
            return clientes
        except Exception as e:
            logging.error(f"Error al leer clientes: {e}")
            return []
            
    @log_performance
    def get_configuracion(self):
        """Lee la configuración institucional y codifica el logo en Base64."""
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM configuracion WHERE id = 1')
            row = cursor.fetchone()
            conn.close()

            conf = dict(row) if row else {}
            logo_path = conf.get('logo_path')

            if logo_path:
                # Si la ruta es absoluta (nueva lógica) o relativa
                if os.path.isabs(logo_path):
                    abs_logo_path = logo_path
                elif logo_path.startswith('../'):
                    abs_logo_path = os.path.join(BASE_DIR, logo_path[3:])
                else:
                    abs_logo_path = os.path.join(BASE_DIR, 'assets', os.path.basename(logo_path))

                if os.path.exists(abs_logo_path):
                    with open(abs_logo_path, 'rb') as img_file:
                        encoded = base64.b64encode(img_file.read()).decode('utf-8')
                        ext = os.path.splitext(abs_logo_path)[1][1:].lower()
                        if ext == 'jpg':
                            ext = 'jpeg'
                        conf['logo_base64'] = f"data:image/{ext};base64,{encoded}"
                else:
                    # Intentar buscar en WRITABLE_DIR si no se encontró (migración)
                    alt_path = os.path.join(WRITABLE_DIR, 'assets', os.path.basename(logo_path))
                    if os.path.exists(alt_path):
                        with open(alt_path, 'rb') as img_file:
                            encoded = base64.b64encode(img_file.read()).decode('utf-8')
                            conf['logo_base64'] = f"data:image/png;base64,{encoded}" # Fallback a png
                        conf['logo_path'] = alt_path # Auto-corrección
                    else:
                        conf['logo_base64'] = None
            else:
                conf['logo_base64'] = None

            return conf
        except Exception as e:
            logging.error(f"Error al leer configuracion: {e}")
            return {}

    def quitar_logo(self):
        """Elimina el logo de la base de datos."""
        try:
            conn = _get_db_connection()
            conn.execute("UPDATE configuracion SET logo_path = NULL WHERE id = 1")
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al quitar logo: {e}")
            return str(e)

    def guardar_configuracion(self, nombre, lema, pin=None, exequatur=None, tomo=None, folio_ini=None, limite=None):
        """Guarda la identidad institucional, metadatos de auditoría y PIN de seguridad."""
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            
            # Actualizar campos básicos
            cursor.execute('''
                UPDATE configuracion 
                SET nombre_despacho = ?, lema_legal = ? 
                WHERE id = 1
            ''', (nombre, lema))
            
            # Actualizar metadatos de auditoría
            if exequatur is not None:
                cursor.execute("UPDATE configuracion SET exequatur = ? WHERE id = 1", (exequatur,))
            if tomo is not None:
                cursor.execute("UPDATE configuracion SET tomo_actual = ? WHERE id = 1", (tomo,))
            if folio_ini is not None:
                cursor.execute("UPDATE configuracion SET folio_inicial_tomo = ? WHERE id = 1", (folio_ini,))
            if limite is not None:
                cursor.execute("UPDATE configuracion SET limite_tomo = ? WHERE id = 1", (limite,))

            # Actualizar PIN si se proporciona uno válido
            if pin and len(pin) >= 4:
                cursor.execute("UPDATE configuracion SET pin = ? WHERE id = 1", (pin,))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al guardar configuracion: {e}")
            return str(e)

    def verificar_pin(self, pin_ingresado):
        """Verifica el PIN de administrador. Si no hay PIN configurado, permite todo."""
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT pin FROM configuracion WHERE id=1")
            row = cursor.fetchone()
            conn.close()
            pin_guardado = row['pin'] if row and row['pin'] else ''
            if not pin_guardado:
                return True
            return pin_ingresado == pin_guardado
        except Exception:
            return True  # Fallo seguro: si hay error, no bloquear

    def reset_pin_admin(self, codigo_reset):
        """Resetea el PIN usando el código maestro del sistema."""
        CODIGO_MAESTRO = "JURIS2026RESET"
        if codigo_reset.strip().upper() != CODIGO_MAESTRO:
            return "Código de restablecimiento incorrecto."
        try:
            conn = _get_db_connection()
            conn.execute("UPDATE configuracion SET pin='' WHERE id=1")
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            return str(e)

    # ---- CRUD CLIENTES ----
    def editar_cliente(self, id_cliente, nombre, dni, telefono, email, direccion):
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE clientes SET nombre_completo=?, dni=?, telefono=?, correo=?, direccion=?
                WHERE id_cliente=?
            ''', (nombre, dni, telefono, email, direccion, id_cliente))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al editar cliente: {e}")
            return str(e)

    def eliminar_cliente(self, id_cliente):
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute("DELETE FROM clientes WHERE id_cliente=?", (id_cliente,))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al eliminar cliente: {e}")
            return str(e)

    # ---- CRUD EXPEDIENTES ----
    def editar_expediente(self, id_expediente, materia, juzgado):
        try:
            conn = _get_db_connection()
            conn.execute(
                'UPDATE expedientes SET materia=?, juzgado=? WHERE id_expediente=?',
                (materia, juzgado, id_expediente)
            )
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al editar expediente: {e}")
            return str(e)

    def eliminar_expediente(self, id_expediente):
        try:
            conn = _get_db_connection()
            conn.execute("DELETE FROM actuaciones WHERE id_expediente=?", (id_expediente,))
            conn.execute("DELETE FROM expedientes WHERE id_expediente=?", (id_expediente,))
            conn.commit()
            conn.close()
            logging.info(f"Expediente eliminado: {id_expediente}")
            return True
        except Exception as e:
            logging.error(f"Error al eliminar expediente: {e}")
            return str(e)

    # ---- CRUD AUDIENCIAS ----
    def editar_audiencia(self, id_actuacion, tipo, fecha_hora, observaciones, monto):
        try:
            conn = _get_db_connection()
            conn.execute('''
                UPDATE actuaciones SET tipo_audiencia=?, fecha_hora=?, observaciones=?, monto_involucrado=?
                WHERE id_actuacion=?
            ''', (tipo, fecha_hora, observaciones, monto, id_actuacion))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al editar audiencia: {e}")
            return str(e)

    def eliminar_audiencia(self, id_actuacion):
        try:
            conn = _get_db_connection()
            conn.execute("DELETE FROM actuaciones WHERE id_actuacion=?", (id_actuacion,))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al eliminar audiencia: {e}")
            return str(e)

    # ---- CRUD PROTOCOLO ----
    def eliminar_instrumento(self, id_instrumento):
        try:
            conn = _get_db_connection()
            conn.execute("DELETE FROM protocolo WHERE id_instrumento=?", (id_instrumento,))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error al eliminar instrumento: {e}")
            return str(e)

    def seleccionar_logo(self):
        """Abre un diálogo para seleccionar una imagen, la guarda y retorna Base64."""
        file_types = ('Imágenes (*.png;*.jpg;*.jpeg)', 'Todos los archivos (*.*)')
        result = webview.windows[0].create_file_dialog(
            webview.OPEN_DIALOG, allow_multiple=False, file_types=file_types
        )
        if not result:
            return None

        source_path = result[0]
        # Guardar SIEMPRE en el directorio escribible del usuario
        assets_dir = os.path.join(WRITABLE_DIR, 'assets')
        os.makedirs(assets_dir, exist_ok=True)

        ext = os.path.splitext(source_path)[1]
        dest_path = os.path.join(assets_dir, f'custom_logo{ext}')
        shutil.copy2(source_path, dest_path)

        conn = _get_db_connection()
        # Guardamos la ruta absoluta en la DB para que sea inequívoca
        conn.execute(
            "UPDATE configuracion SET logo_path = ? WHERE id = 1",
            (dest_path,)
        )
        conn.commit()
        conn.close()

        with open(dest_path, 'rb') as img_file:
            encoded = base64.b64encode(img_file.read()).decode('utf-8')
            mime = ext[1:].lower()
            if mime == 'jpg':
                mime = 'jpeg'
            return f"data:image/{mime};base64,{encoded}"

    def get_ai_logo(self):
        """Retorna el logo del agente IA en base64 para el panel derecho."""
        path = os.path.join(BASE_DIR, 'assets', 'ia agente.png')
        if not os.path.exists(path):
            return None
        try:
            with open(path, 'rb') as img_file:
                encoded = base64.b64encode(img_file.read()).decode('utf-8')
                return f"data:image/png;base64,{encoded}"
        except Exception as e:
            logging.error(f"Error al cargar logo IA: {e}")
            return None

    def seleccionar_archivo_incidentes(self):
        """Abre un diálogo nativo para enlazar un archivo JSON de incidentes externo."""
        file_types = ('Archivos JSON (*.json)', 'Todos los archivos (*.*)')
        result = webview.windows[0].create_file_dialog(
            webview.OPEN_DIALOG, allow_multiple=False, file_types=file_types
        )
        if not result:
            return None

        source_path = result[0]
        conn = _get_db_connection()
        conn.execute("UPDATE configuracion SET incidentes_path = ? WHERE id = 1", (source_path,))
        conn.commit()
        conn.close()
        return source_path

    def get_audiencias(self):
        """Consulta el calendario de audiencias programadas con todos los detalles del cliente y caso."""
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                SELECT a.id_actuacion, a.fecha_hora, a.tipo_audiencia, a.observaciones, 
                       a.monto_involucrado,
                       e.id_expediente, e.juzgado, e.materia,
                       c.nombre_completo AS cliente_nombre
                FROM actuaciones a
                JOIN expedientes e ON a.id_expediente = e.id_expediente
                JOIN clientes c ON e.id_cliente = c.id_cliente
                ORDER BY a.fecha_hora ASC
            ''')
            audiencias = [dict(row) for row in cursor.fetchall()]
            conn.close()
            return audiencias
        except Exception as e:
            logging.error(f"Error al leer audiencias: {e}")
            return []

    def get_lista_plantillas(self):
        """Escanea la carpeta de plantillas y devuelve una lista para el UI."""
        plantillas = []
        base_path = os.path.join(BASE_DIR, 'plantillas')
        if not os.path.exists(base_path): return []
        for root, dirs, files in os.walk(base_path):
            for file in files:
                if file.endswith('.md'):
                    ruta = os.path.relpath(os.path.join(root, file), base_path)
                    plantillas.append({'nombre': file.replace('.md', '').replace('_', ' ').title(), 'ruta': ruta})
        return sorted(plantillas, key=lambda x: x['nombre'])

    def cargar_plantilla(self, ruta_relativa, id_expediente):
        """Carga el texto de la plantilla e inyecta los datos del cliente/caso activo."""
        plantilla_path = os.path.join(BASE_DIR, 'plantillas', ruta_relativa)
        if not os.path.exists(plantilla_path):
            return "Error: Plantilla no encontrada."

        try:
            with open(plantilla_path, 'r', encoding='utf-8') as f:
                texto = f.read()

            conn = _get_db_connection()
            cursor = conn.cursor()
            
            # Buscar datos del expediente y cliente
            cursor.execute('''
                SELECT e.juzgado, c.nombre_completo, c.rtn, c.dni, c.telefono, c.direccion 
                FROM expedientes e JOIN clientes c ON e.id_cliente = c.id_cliente 
                WHERE e.id_expediente = ?
            ''', (id_expediente,))
            row = cursor.fetchone()
            
            if row:
                texto = texto.replace("[NOMBRE_CLIENTE]", row['nombre_completo'])
                texto = texto.replace("[DNI_CLIENTE]", row['dni'] or "")
                texto = texto.replace("[TELEFONO]", row['telefono'] or "")
                texto = texto.replace("[DOMICILIO_CLIENTE]", row['direccion'] or "")
                if row['juzgado']:
                    texto = texto.replace("[JUZGADO CORRESPONDIENTE]", row['juzgado'])
            conn.close()
            return texto
        except Exception as e:
            return f"Error al cargar la plantilla: {e}"

    def guardar_documento(self, id_expediente, contenido):
        """Genera un .docx en papel legal, Times New Roman 14pt, 1.5 espaciado y abre diálogo de guardado."""
        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re

            doc = Document()

            # --- Papel Legal: 8.5" × 14" con márgenes de 2.5 cm ---
            section = doc.sections[0]
            section.page_width  = Inches(8.5)
            section.page_height = Inches(14)
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

            # --- Estilo de párrafo por defecto ---
            style = doc.styles['Normal']
            font  = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14)

            # Asegurar que el XML del estilo use Times New Roman en todos los contextos
            rPr = style.element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'),    'Times New Roman')
            rFonts.set(qn('w:hAnsi'),   'Times New Roman')
            rFonts.set(qn('w:cs'),       'Times New Roman')
            rPr.insert(0, rFonts)

            def _aplicar_formato_parrafo(p):
                """Aplica Times New Roman 14pt y espaciado 1.5 a un párrafo."""
                pPr = p._p.get_or_add_pPr()
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:line'),      '360')   # 1.5 × 240 = 360 twips
                spacing.set(qn('w:lineRule'), 'auto')
                spacing.set(qn('w:before'),   '0')
                spacing.set(qn('w:after'),    '120')
                pPr.append(spacing)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

            def _parrafo_nuevo(texto, negrita=False, centrado=False):
                p = doc.add_paragraph()
                run = p.add_run(texto)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = negrita
                if centrado:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _aplicar_formato_parrafo(p)
                return p

            # --- Convertir contenido (markdown simple) a párrafos Word ---
            lineas = contenido.splitlines()
            for linea in lineas:
                stripped = linea.strip()

                if not stripped:
                    # Párrafo vacío como separador
                    p = doc.add_paragraph()
                    _aplicar_formato_parrafo(p)
                    continue

                # Encabezados Markdown → negrita centrada
                if stripped.startswith('### '):
                    _parrafo_nuevo(stripped[4:], negrita=True, centrado=True)
                elif stripped.startswith('## '):
                    _parrafo_nuevo(stripped[3:], negrita=True, centrado=True)
                elif stripped.startswith('# '):
                    _parrafo_nuevo(stripped[2:], negrita=True, centrado=True)
                # Listas con guion o asterisco
                elif re.match(r'^[-*]\s+', stripped):
                    texto_item = re.sub(r'^[-*]\s+', '', stripped)
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(texto_item)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    _aplicar_formato_parrafo(p)
                # Separadores horizontales
                elif stripped in ('---', '***', '___'):
                    p = doc.add_paragraph()
                    pPr = p._p.get_or_add_pPr()
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'),   'single')
                    bottom.set(qn('w:sz'),    '6')
                    bottom.set(qn('w:space'), '1')
                    bottom.set(qn('w:color'), '000000')
                    pBdr.append(bottom)
                    pPr.append(pBdr)
                    _aplicar_formato_parrafo(p)
                else:
                    # Párrafo normal — procesar **negrita** inline
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    segmentos = re.split(r'\*\*(.+?)\*\*', stripped)
                    for i, seg in enumerate(segmentos):
                        if not seg:
                            continue
                        run = p.add_run(seg)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        run.bold = (i % 2 == 1)  # impares son negrita
                    _aplicar_formato_parrafo(p)

            # --- Diálogo nativo de guardado ---
            default_filename = f"{id_expediente}_Documento_Legal.docx"
            file_types = ('Documento Word (*.docx)',)

            result = webview.windows[0].create_file_dialog(
                webview.FileDialog.SAVE,
                save_filename=default_filename,
                file_types=file_types
            )

            if result and len(result) > 0:
                save_path = result[0]
                if not save_path.lower().endswith('.docx'):
                    save_path += '.docx'
                doc.save(save_path)
                return True
            return False
        except Exception as e:
            logging.error(f"Error al generar docx: {e}")
            return str(e)

    def actualizar_estado_expediente(self, id_expediente, nuevo_estado):
        """Actualiza el estado de un expediente (Activo, En Proceso, Cerrado)."""
        estados_validos = ['Activo', 'En Proceso', 'Cerrado']
        if nuevo_estado not in estados_validos:
            return f"Estado inválido. Use: {', '.join(estados_validos)}"
        try:
            conn = _get_db_connection()
            cursor = conn.execute(
                "UPDATE expedientes SET estado = ? WHERE id_expediente = ?",
                (nuevo_estado, id_expediente)
            )
            if cursor.rowcount == 0:
                conn.close()
                return f"No se encontró el expediente '{id_expediente}'."
            conn.commit()
            conn.close()
            logging.info(f"Expediente {id_expediente} -> {nuevo_estado}")
            return True
        except Exception as e:
            logging.error(f"Error actualizando estado: {e}")
            return str(e)

    def registrar_cliente(self, nombre, dni, telefono, email, direccion):
        """Registra un nuevo cliente en el CRM."""
        try:
            conn = _get_db_connection()
            conn.execute('''
                INSERT INTO clientes (nombre_completo, dni, telefono, correo, direccion)
                VALUES (?, ?, ?, ?, ?)
            ''', (nombre, dni, telefono, email, direccion))
            conn.commit()
            conn.close()
            return True
        except sqlite3.IntegrityError:
            return f"Ya existe un cliente registrado con el DNI/RTN '{dni}'. Verifique el número o busque el cliente existente en el directorio."
        except Exception as e:
            logging.error(f"Error registrando cliente: {e}")
            return str(e)

    def get_siguiente_id_expediente(self, prefijo_materia):
        """Genera el siguiente ID secuencial para un expediente según la materia."""
        import datetime
        anio = datetime.datetime.now().year
        patron = f"{anio}-{prefijo_materia}-"
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id_expediente FROM expedientes WHERE id_expediente LIKE ?",
                (f"{patron}%",)
            )
            existentes = [row['id_expediente'] for row in cursor.fetchall()]
            conn.close()
            max_num = 0
            for eid in existentes:
                try:
                    num = int(eid.split('-')[-1])
                    if num > max_num:
                        max_num = num
                except ValueError:
                    pass
            siguiente = max_num + 1
            return f"{patron}{str(siguiente).zfill(3)}"
        except Exception as e:
            logging.error(f"Error generando ID: {e}")
            return f"{patron}001"

    def registrar_expediente(self, id_expediente, id_cliente, materia, juzgado, cuantia):
        """Registra un nuevo expediente vinculado a un cliente."""
        try:
            conn = _get_db_connection()
            conn.execute('''
                INSERT INTO expedientes (id_expediente, id_cliente, materia, juzgado, estado)
                VALUES (?, ?, ?, ?, 'Activo')
            ''', (id_expediente, id_cliente, materia, juzgado))
            if cuantia and cuantia > 0:
                conn.execute('''
                    INSERT INTO actuaciones (id_expediente, tipo_audiencia, fecha_hora, observaciones, monto_involucrado)
                    VALUES (?, 'Cuantía / Honorarios Pactados', DATE('now'), 'Registro inicial de honorarios', ?)
                ''', (id_expediente, cuantia))
            conn.commit()
            conn.close()
            return True
        except sqlite3.IntegrityError:
            return f"El ID de Expediente '{id_expediente}' ya existe. Use un identificador diferente (ej: 2026-CIV-004)."
        except Exception as e:
            logging.error(f"Error registrando expediente: {e}")
            return str(e)

    def registrar_audiencia(self, id_expediente, tipo_audiencia, fecha_hora, observaciones, monto):
        """Registra una nueva audiencia en el calendario con su cuantía/honorarios."""
        try:
            conn = _get_db_connection()
            conn.execute('''
                INSERT INTO actuaciones (id_expediente, tipo_audiencia, fecha_hora, observaciones, monto_involucrado)
                VALUES (?, ?, ?, ?, ?)
            ''', (id_expediente, tipo_audiencia, fecha_hora, observaciones, monto))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error registrando audiencia: {e}")
            return str(e)

    def get_honorarios_totales(self):
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT SUM(monto_involucrado) FROM actuaciones")
            total = cursor.fetchone()[0]
            conn.close()
            return total if total else 0.0
        except Exception as e:
            logging.error(f"Error al calcular honorarios: {e}")
            return 0.0

    def get_protocolo(self):
        try:
            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                SELECT p.id_instrumento, p.numero_instrumento, p.fecha_otorgamiento, 
                       p.naturaleza_acto, p.folio_inicio, p.folio_fin,
                       c.nombre_completo AS cliente_nombre
                FROM protocolo p
                JOIN clientes c ON p.id_cliente = c.id_cliente
                ORDER BY p.numero_instrumento DESC
            ''')
            registros = [dict(row) for row in cursor.fetchall()]
            conn.close()
            return registros
        except Exception as e:
            logging.error(f"Error protocolo: {e}")
            return []

    @log_performance
    def registrar_instrumento(self, id_cliente, naturaleza_acto, cantidad_folios, fecha_solicitada=None):
        """
        Registra un instrumento bajo reglas estrictas de auditoría notarial (ACID).
        Garantiza integridad cronológica, numérica y límite de folios por tomo.
        """
        import datetime
        if not fecha_solicitada:
            fecha_solicitada = datetime.datetime.now().strftime('%Y-%m-%d')

        conn = _get_db_connection()
        try:
            cursor = conn.cursor()
            
            # PASO A: Recuperación de Estado y Metadatos de Control
            cursor.execute("SELECT tomo_actual, limite_tomo, exequatur FROM configuracion WHERE id = 1")
            config = cursor.fetchone()
            if not config:
                raise Exception("Configuración institucional no encontrada.")
                
            tomo_activo = config['tomo_actual']
            limite = config['limite_tomo']

            # Obtener el último instrumento para validar secuencia y cronología
            cursor.execute("""
                SELECT numero_instrumento, fecha_otorgamiento, folio_fin 
                FROM protocolo 
                ORDER BY numero_instrumento DESC LIMIT 1
            """)
            ultimo_registro = cursor.fetchone()

            # PASO B: Lógica de Asignación Automática (Fallo Cero)
            nuevo_numero = (ultimo_registro['numero_instrumento'] + 1) if ultimo_registro else 1
            folio_inicio = (ultimo_registro['folio_fin'] + 1) if ultimo_registro else 1
            
            # Si es el primer instrumento del tomo (o el primero de la historia), 
            # verificar si debemos usar el folio inicial configurado
            if nuevo_numero == 1:
                cursor.execute("SELECT folio_inicial_tomo FROM configuracion WHERE id = 1")
                f_init = cursor.fetchone()
                if f_init and f_init['folio_inicial_tomo'] > 1:
                    folio_inicio = f_init['folio_inicial_tomo']

            folio_fin = folio_inicio + int(cantidad_folios) - 1

            # PASO C: Validaciones de Negocio (Pre-check)
            
            # 1. Validación Cronológica (Evitar inserción retroactiva)
            if ultimo_registro and fecha_solicitada < ultimo_registro['fecha_otorgamiento']:
                raise Exception(f"Inconsistencia Temporal: La fecha {fecha_solicitada} es anterior al último instrumento registrado ({ultimo_registro['fecha_otorgamiento']}). El protocolo debe ser estrictamente cronológico.")

            # 2. Regla de los 200 (Límite de Tomo)
            if folio_fin > limite:
                raise Exception(f"BLOQUEO DE AUDITORÍA: El instrumento solicitado finalizaría en el folio {folio_fin}, excediendo el límite legal de {limite} folios para el Tomo {tomo_activo}. Debe proceder al Cierre de Tomo y Apertura de uno nuevo.")

            # PASO D: Ejecución de la Transacción
            cursor.execute('''
                INSERT INTO protocolo (
                    numero_instrumento, fecha_otorgamiento, id_cliente, 
                    naturaleza_acto, folio_inicio, folio_fin, numero_tomo
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (nuevo_numero, fecha_solicitada, id_cliente, naturaleza_acto, folio_inicio, folio_fin, tomo_activo))

            conn.commit()
            logging.info(f"Instrumento {nuevo_numero} registrado exitosamente (Tomo {tomo_activo}, Folios {folio_inicio}-{folio_fin}).")
            return True

        except Exception as e:
            conn.rollback()
            logging.error(f"Error de validación notarial: {e}")
            return str(e)
        finally:
            conn.close()
            
    def generar_indice_notarial(self):
        try:
            # Obtener registros y configuración para el encabezado
            registros = self.get_protocolo()
            if not registros:
                return "No hay instrumentos en el protocolo para generar el índice."

            conn = _get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT nombre_institucion, lema_legal, exequatur, tomo_actual FROM configuracion WHERE id = 1")
            config = cursor.fetchone()
            conn.close()

            nombre_notario = config['nombre_institucion'] if config else "Notario no configurado"
            exequatur = config['exequatur'] if config else "S/D"
            tomo = config['tomo_actual'] if config else "1"

            lineas = [
                f"# ÍNDICE DEL PROTOCOLO NOTARIAL",
                f"**Notario:** {nombre_notario}",
                f"**Exequátur:** {exequatur} | **Tomo:** {tomo}\n",
                "| No. Instrumento | Fecha | Otorgante / Cliente | Naturaleza del Acto | Folios |",
                "|---|---|---|---|---|",
            ]
            # Registros en orden cronológico para el índice (del 1 al N)
            for r in sorted(registros, key=lambda x: x['numero_instrumento']):
                lineas.append(
                    f"| {r['numero_instrumento']} | {r['fecha_otorgamiento']} | "
                    f"{r['cliente_nombre']} | {r['naturaleza_acto']} | "
                    f"{r['folio_inicio']} al {r['folio_fin']} |"
                )
            
            # Pie de página de auditoría
            lineas.append(f"\n*Generado automáticamente por JURIS-GESTIÓN-PRO el {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}*")
            
            contenido = "\n".join(lineas)

            result = webview.windows[0].create_file_dialog(
                webview.FileDialog.SAVE,
                save_filename=f"Indice_Notarial_Tomo_{tomo}.md",
                file_types=('Documento Markdown (*.md)', 'Todos los archivos (*.*)')
            )

            if result:
                with open(result[0], 'w', encoding='utf-8') as f:
                    f.write(contenido)
                return True
            return False
        except Exception as e:
            logging.error(f"Error al generar indice notarial: {e}")
            return str(e)

    def procesar_mensaje_ia(self, mensaje, id_expediente=None):
        """Procesa los mensajes del usuario interactuando con las plantillas o la API de IA."""
        mensaje_lower = mensaje.lower()

        if "incidente" in mensaje_lower:
            try:
                conn = _get_db_connection()
                ruta_incidentes = self._get_ruta_incidentes(conn.cursor())
                conn.close()
                with open(ruta_incidentes, 'r', encoding='utf-8') as f:
                    incidentes = json.load(f)
                resumen = "<strong>Catálogo de Incidentes Procesales:</strong><br>"
                for i in incidentes:
                    resumen += f"&#x1F538; <b>{i['nombre']}</b> ({i['materia']})<br>"
                return resumen
            except Exception:
                return "Error al leer la base de incidentes desde el archivo externo configurado."

        if "demanda" in mensaje_lower or "redacta" in mensaje_lower or "autocompleta" in mensaje_lower:
            plantilla_path = os.path.join(BASE_DIR, 'plantillas', 'demanda_abreviado.md')
            if not os.path.exists(plantilla_path):
                return "No encontré la plantilla oficial en la biblioteca."
            try:
                with open(plantilla_path, 'r', encoding='utf-8') as f:
                    texto = f.read()

                # Inyectar datos reales desde configuración institucional
                conf = self.get_configuracion()
                nombre_despacho = conf.get('nombre_despacho', '[NOMBRE_ABOGADO]')
                texto = texto.replace("[NOMBRE_ABOGADO]", nombre_despacho)

                # Si hay expediente activo, inyectar datos del cliente/caso
                tiene_expediente = False
                if id_expediente:
                    try:
                        conn = _get_db_connection()
                        cursor = conn.cursor()
                        cursor.execute('''
                            SELECT e.juzgado, e.materia, c.nombre_completo, c.dni, c.telefono, c.direccion 
                            FROM expedientes e JOIN clientes c ON e.id_cliente = c.id_cliente 
                            WHERE e.id_expediente = ?
                        ''', (id_expediente,))
                        row = cursor.fetchone()
                        conn.close()
                        if row:
                            tiene_expediente = True
                            texto = texto.replace("[NOMBRE_CLIENTE]", row['nombre_completo'] or "")
                            texto = texto.replace("[DNI_CLIENTE]", row['dni'] or "")
                            texto = texto.replace("[TELEFONO]", row['telefono'] or "")
                            texto = texto.replace("[DOMICILIO_CLIENTE]", row['direccion'] or "")
                            if row['juzgado']:
                                texto = texto.replace("[DEPARTAMENTO]", row['juzgado'])
                    except Exception as e:
                        logging.warning(f"Error inyectando datos del expediente: {e}")

                if not tiene_expediente:
                    texto = texto.replace("[NOMBRE_CLIENTE]", "[Cliente por definir]")

                texto = texto.replace("[NUMERO_COLEGIO]", "[No. Colegio]")

                vista_previa = texto[:350] + "..."
                return {
                    "tipo": "borrador",
                    "texto_completo": texto,
                    "vista_previa": vista_previa,
                    "tiene_expediente": tiene_expediente
                }
            except Exception:
                return "Error al procesar la plantilla."

        return (
            "JURIS-AI Activo. Para análisis conversacional avanzado o resúmenes de jurisprudencia, "
            "asegúrate de activar la conexión API en la Configuración. "
            "Por ahora, puedo autocompletar 'demandas' y consultar 'incidentes'."
        )

    def generar_backup_manual(self):
        """Genera un respaldo manualmente desde la interfaz."""
        try:
            from src.seguridad import ejecutar_escudo_datos
            exito = ejecutar_escudo_datos()
            if exito:
                return "Respaldo de seguridad (Escudo de Datos) generado exitosamente."
            return "Hubo un error al generar el respaldo."
        except Exception as e:
            logging.error(f"Error backup manual: {e}")
            return str(e)


def al_cerrar():
    """Evento que se dispara al cerrar la ventana. Activa el Escudo de Datos."""
    from src.seguridad import ejecutar_escudo_datos
    logging.info("Iniciando secuencia de cierre...")
    ejecutar_escudo_datos()

if __name__ == '__main__':
    logging.info("Iniciando JURIS-GESTIÓN-PRO (Aplicación de Escritorio Nativa)...")
    
    api = JurisAPI()
    
    # Crear ventana de escritorio nativa sin barra de navegación de navegador
    ventana = webview.create_window(
        title='JURIS-GESTIÓN-PRO | Control de Casos y Protocolo', 
        url=UI_PATH, 
        js_api=api,
        width=1280, 
        height=800,
        min_size=(1024, 768),
        maximized=True,
        background_color='#f3f3f3' # Fondo claro nativo por defecto
    )
    
    # Vincular el evento de cierre al Escudo de Datos
    ventana.events.closed += al_cerrar
    
    # Iniciar la aplicación
    webview.start(debug=False)
